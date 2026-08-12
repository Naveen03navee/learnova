import io
from docx import Document
from docx.shared import Pt, Inches
from app.models.paper import QuestionPaper, PaperStatus

def export_question_paper_docx(paper: QuestionPaper) -> io.BytesIO:
    if paper.status != PaperStatus.APPROVED:
        raise ValueError("Only APPROVED papers can be exported.")
        
    doc = Document()
    
    # Title and Metadata
    title = doc.add_heading(paper.title, level=1)
    title.alignment = 1 # Center
    
    doc.add_paragraph(f"Date: __________________\nStudent Name: __________________\n", style="Normal")
    
    # Sort items by order_index
    items = sorted(paper.items, key=lambda x: x.order_index)
    
    # Group by section
    sections = {}
    for item in items:
        if item.section_name not in sections:
            sections[item.section_name] = []
        sections[item.section_name].append(item)
        
    question_number = 1
    
    for section_name, section_items in sections.items():
        doc.add_heading(section_name, level=2)
        
        for item in section_items:
            # Add question text and marks
            marks = item.marks_override if item.marks_override is not None else item.marks_snapshot
            p = doc.add_paragraph(f"{question_number}. {item.question_text_snapshot} ")
            p.add_run(f"[{marks} Marks]").italic = True
            
            # Add options if present (MCQ)
            options = item.content_snapshot.get("options")
            if options:
                for opt in options:
                    doc.add_paragraph(f"   {opt['id']}. {opt['text']}", style="Normal")
                    
            # Add space for answer if not MCQ
            if not options:
                doc.add_paragraph("\n\n")
                
            question_number += 1
            
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
