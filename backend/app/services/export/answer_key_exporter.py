import io
from docx import Document
from fpdf import FPDF
from app.models.paper import QuestionPaper, PaperStatus

def export_answer_key_docx(paper: QuestionPaper) -> io.BytesIO:
    if paper.status != PaperStatus.APPROVED:
        raise ValueError("Only APPROVED papers can be exported.")
        
    doc = Document()
    
    # Title and Metadata
    title = doc.add_heading(f"{paper.title} - Answer Key", level=1)
    title.alignment = 1 # Center
    
    # Sort items by order_index
    items = sorted(paper.items, key=lambda x: x.order_index)
    
    # Group by section
    sections = {}
    for item in items:
        if item.section_name not in sections:
            sections[item.section_name] = []
        sections[item.section_name].append(item)
        
    question_number = 1
    
    for section_name, section_items in sections.items():
        doc.add_heading(section_name, level=2)
        
        for item in section_items:
            marks = item.marks_override if item.marks_override is not None else item.marks_snapshot
            doc.add_paragraph(f"{question_number}. ", style="Heading 3")
            
            correct_answer = item.content_snapshot.get("correct_answer", "N/A")
            explanation = item.content_snapshot.get("explanation", "N/A")
            
            p1 = doc.add_paragraph()
            p1.add_run("Correct Answer: ").bold = True
            p1.add_run(str(correct_answer))
            
            p2 = doc.add_paragraph()
            p2.add_run("Explanation: ").bold = True
            p2.add_run(str(explanation))
            
            p3 = doc.add_paragraph()
            p3.add_run("Marks: ").bold = True
            p3.add_run(str(marks))
            
            doc.add_paragraph("\n")
            question_number += 1
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _clean_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "--",
        "\u2026": "...",
        "\u00a0": " ",
        "\u2264": "<=", "\u2265": ">=",
        "\u2260": "!=", "\u00b1": "+/-",
        "\u00d7": "*", "\u00f7": "/",
        "\u03bc": "u", "\u03a9": "Ohm",
        "\u03c0": "pi", "\u03b8": "theta",
        "\u03bb": "lambda", "\u03b1": "alpha", "\u03b2": "beta", "\u03b3": "gamma"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")

def export_answer_key_pdf(paper: QuestionPaper) -> io.BytesIO:
    if paper.status != PaperStatus.APPROVED:
        raise ValueError("Only APPROVED papers can be exported.")
        
    is_college_exam = paper.exam and 'college' in paper.exam.name.lower()
        
    pdf = FPDF()
    pdf.add_page()
    
    # Fonts
    pdf.set_font("Helvetica", "B", 16)
    
    # Title
    title = _clean_text(f"{paper.title} - Answer Key")
    try:
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align='C')
    except:
        pdf.cell(0, 10, title, ln=True, align='C')
    pdf.ln(5)
    
    # Sort items by order_index
    items = sorted(paper.items, key=lambda x: x.order_index)
    
    # Group by section
    sections = {}
    for item in items:
        if item.section_name not in sections:
            sections[item.section_name] = []
        sections[item.section_name].append(item)
        
    question_number = 1
    
    for section_name, section_items in sections.items():
        pdf.set_font("Helvetica", "B", 14)
        s_name = _clean_text(section_name)
        try:
            pdf.cell(0, 10, s_name, new_x="LMARGIN", new_y="NEXT")
        except:
            pdf.cell(0, 10, s_name, ln=True)
        pdf.ln(2)
        
        if is_college_exam:
            for item in section_items:
                marks = item.marks_override if item.marks_override is not None else item.marks_snapshot
                correct_answer = _clean_text(str(item.content_snapshot.get("correct_answer", "N/A")))
                explanation = _clean_text(str(item.content_snapshot.get("explanation", "N/A")))
                
                pdf.set_font("Helvetica", "B", 12)
                try:
                    pdf.cell(0, 6, f"{question_number}. ", new_x="LMARGIN", new_y="NEXT")
                except:
                    pdf.cell(0, 6, f"{question_number}. ", ln=True)
                
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(35, 6, "Correct Answer: ")
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 6, correct_answer)
                pdf.set_x(pdf.l_margin)
                
                if explanation and explanation != "N/A":
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.cell(30, 6, "Explanation: ")
                    pdf.set_font("Helvetica", "", 11)
                    pdf.multi_cell(0, 6, explanation)
                    pdf.set_x(pdf.l_margin)
                
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(15, 6, "Marks: ")
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 6, str(marks))
                pdf.set_x(pdf.l_margin)
                
                pdf.ln(4)
                question_number += 1
        else:
            cols = 5
            usable_width = pdf.w - 2 * pdf.l_margin
            col_width = usable_width / cols
            q_width = col_width * 0.4
            a_width = col_width * 0.6
            
            # Header
            pdf.set_font("Helvetica", "B", 10)
            for _ in range(cols):
                pdf.cell(q_width, 8, "Q", border=1, align='C')
                pdf.cell(a_width, 8, "Ans", border=1, align='C')
            pdf.ln(8)
            
            pdf.set_font("Helvetica", "", 10)
            items_per_row = cols
            
            for i in range(0, len(section_items), items_per_row):
                row_items = section_items[i:i+items_per_row]
                for j, item in enumerate(row_items):
                    correct_answer = _clean_text(str(item.content_snapshot.get("correct_answer", "N/A")))
                    pdf.cell(q_width, 8, str(question_number + j), border=1, align='C')
                    pdf.cell(a_width, 8, correct_answer, border=1, align='C')
                
                # If the last row isn't full, pad with empty cells
                if len(row_items) < cols:
                    for _ in range(cols - len(row_items)):
                        pdf.cell(q_width, 8, "", border=1, align='C')
                        pdf.cell(a_width, 8, "", border=1, align='C')
                        
                question_number += len(row_items)
                pdf.ln(8)
            pdf.ln(4)
            
    buffer = io.BytesIO()
    try:
        buffer.write(pdf.output())
    except TypeError:
        buffer.write(pdf.output(dest='S').encode('latin1'))
    buffer.seek(0)
    return buffer


