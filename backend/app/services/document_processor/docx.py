import io
import docx

def extract_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Extracts paragraphs and basic table content.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    
    content = []
    
    # Simple extraction of all paragraphs in order
    # (python-docx doesn't perfectly order paragraphs vs tables if they are interleaved in a complex way, 
    # but for simple educational notes this is usually sufficient).
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
            
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    # Clean newlines inside table cells to spaces to keep row coherent
                    row_data.append(text.replace('\n', ' '))
            if row_data:
                content.append(" | ".join(row_data))
                
    return "\n\n".join(content)
