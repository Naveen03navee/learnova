import unittest
import sys
import os
from uuid import uuid4
from unittest.mock import MagicMock
import io
from docx import Document

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app.models.paper import QuestionPaper, QuestionPaperItem, PaperStatus
from app.services.export.docx_exporter import export_question_paper_docx
from app.services.export.answer_key_exporter import export_answer_key_docx

class TestExports(unittest.TestCase):
    
    def setUp(self):
        self.paper = MagicMock(spec=QuestionPaper)
        self.paper.id = uuid4()
        self.paper.title = "Test Paper"
        self.paper.status = PaperStatus.APPROVED
        
        self.item1 = MagicMock(spec=QuestionPaperItem)
        self.item1.section_name = "Section A"
        self.item1.order_index = 0
        self.item1.marks_snapshot = 5
        self.item1.marks_override = None
        self.item1.question_text_snapshot = "Snapshot Question Text 1"
        self.item1.content_snapshot = {
            "correct_answer": "Snapshot Answer",
            "explanation": "Snapshot Explanation"
        }
        
        self.paper.items = [self.item1]
        
    def test_export_draft_raises_error(self):
        self.paper.status = PaperStatus.DRAFT
        with self.assertRaises(ValueError) as context:
            export_question_paper_docx(self.paper)
        self.assertIn("Only APPROVED papers can be exported", str(context.exception))
        
        with self.assertRaises(ValueError) as context:
            export_answer_key_docx(self.paper)
        self.assertIn("Only APPROVED papers can be exported", str(context.exception))

    def test_export_question_paper_uses_snapshot(self):
        buffer = export_question_paper_docx(self.paper)
        self.assertIsInstance(buffer, io.BytesIO)
        
        # Parse docx
        doc = Document(buffer)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        self.assertIn("Test Paper", text)
        self.assertIn("Snapshot Question Text 1", text)
        self.assertIn("[5 Marks]", text)
        
    def test_export_answer_key_uses_snapshot(self):
        buffer = export_answer_key_docx(self.paper)
        self.assertIsInstance(buffer, io.BytesIO)
        
        # Parse docx
        doc = Document(buffer)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        self.assertIn("Test Paper - Answer Key", text)
        self.assertIn("Correct Answer: Snapshot Answer", text)
        self.assertIn("Explanation: Snapshot Explanation", text)
        self.assertIn("Marks: 5", text)

if __name__ == '__main__':
    unittest.main()
